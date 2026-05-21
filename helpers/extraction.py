"""PDF and DOCX text extraction for the upload-document RAG ingestion path."""

from __future__ import annotations

import io

import docx
from pypdf import PdfReader


def extract_pdf_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx_text(docx_bytes: bytes) -> str:
    # python-docx's doc.paragraphs SKIPS tables. Fitness plans are commonly
    # laid out as tables, so the extra loop below is what keeps RAG non-empty.
    doc = docx.Document(io.BytesIO(docx_bytes))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
